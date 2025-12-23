"""
Word Document Data Source Connector
Extracts text from .docx files for RAG training.
"""
import os
from typing import List, Dict, Any
from langchain_core.documents import Document
from .base import BaseDataSource

try:
    from docx import Document as DocxDocument
    from docx.table import Table
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARN] python-docx not installed. Word document support disabled.")


class WordSource(BaseDataSource):
    """Extract documents from Word (.docx) files"""
    
    def __init__(self, file_paths: List[str]):
        """
        Initialize Word source with file paths.
        
        Args:
            file_paths: List of paths to .docx files
        """
        self.file_paths = file_paths
        self.documents: List[Document] = []
        
    def get_source_type(self) -> str:
        return "word"
    
    def extract_documents(self) -> List[Document]:
        """
        Extract documents from Word files.
        Preserves paragraph structure and extracts tables as text.
        """
        if not DOCX_AVAILABLE:
            print("[ERROR] python-docx not installed. Run: pip install python-docx")
            return []
            
        all_documents = []
        
        for file_path in self.file_paths:
            if not os.path.exists(file_path):
                print(f"[WARN] Word file not found: {file_path}")
                continue
                
            if not file_path.lower().endswith('.docx'):
                print(f"[WARN] Not a .docx file: {file_path}")
                continue
                
            try:
                docs = self._process_word_file(file_path)
                all_documents.extend(docs)
                print(f"[OK] Extracted {len(docs)} documents from {os.path.basename(file_path)}")
            except Exception as e:
                print(f"[ERROR] Failed to process {file_path}: {e}")
                
        self.documents = all_documents
        return all_documents
    
    def _process_word_file(self, file_path: str) -> List[Document]:
        """Process a single Word file into documents"""
        documents = []
        filename = os.path.basename(file_path)
        
        doc = DocxDocument(file_path)
        
        # Extract all text content
        full_text = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            if table_text:
                full_text.append(f"\n[Table {table_idx + 1}]\n{table_text}")
        
        # Combine into single document or split by sections
        if full_text:
            content = "\n\n".join(full_text)
            
            # Split into chunks if document is large
            chunks = self._split_into_chunks(content, chunk_size=1500, overlap=200)
            
            for chunk_idx, chunk in enumerate(chunks):
                doc_obj = Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "source_type": "word",
                        "chunk_index": chunk_idx,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc_obj)
        
        return documents
    
    def _extract_table_text(self, table: 'Table') -> str:
        """Extract text from a Word table"""
        rows_text = []
        
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        
        return "\n".join(rows_text)
    
    def _split_into_chunks(self, text: str, chunk_size: int = 1500, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end near chunk boundary
                for delim in ['. ', '.\n', '\n\n', '\n']:
                    last_delim = text.rfind(delim, start + chunk_size // 2, end)
                    if last_delim != -1:
                        end = last_delim + len(delim)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    def get_metadata(self) -> Dict[str, Any]:
        """Return metadata about the Word source"""
        return {
            "source_type": "word",
            "file_count": len(self.file_paths),
            "document_count": len(self.documents),
            "files": [os.path.basename(f) for f in self.file_paths]
        }
